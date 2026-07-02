from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DOCX = "Wesley_Simplicio_CV_Data_Evaluations_Engineer.docx"

BLUE = RGBColor(15, 91, 216)
INK = RGBColor(19, 34, 56)
MUTED = RGBColor(86, 101, 121)
LINE = "D8E0EA"
SOFT = "F5F8FC"


def set_page(section):
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D8E0EA", sz="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), sz)
        element.set(qn("w:color"), color)


def set_table_borders(table, color="D8E0EA", sz="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), sz)
        element.set(qn("w:color"), color)


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_run(paragraph, text, *, bold=False, color=INK, size=10.5, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return run


def add_section_title(cell, text):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=4, after=6)
    r = add_run(p, text.upper(), bold=True, color=INK, size=9.4)
    r.font.all_caps = True
    border_p = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), LINE)
    border_p.append(bottom)
    p._p.get_or_add_pPr().append(border_p)
    return p


def add_body(cell, text):
    p = cell.add_paragraph()
    set_paragraph_spacing(p, after=6, line=1.12)
    add_run(p, text, size=10.2, color=INK)
    return p


def add_bullets(cell, items):
    for item in items:
        p = cell.add_paragraph(style="List Bullet")
        set_paragraph_spacing(p, after=4, line=1.1)
        add_run(p, item, size=10.1, color=INK)


def add_exp_entry(cell, company, dates, role, bullets):
    p = cell.add_paragraph()
    set_paragraph_spacing(p, before=4, after=1)
    add_run(p, company, bold=True, size=10.4, color=INK)
    add_run(p, " | ", size=10.0, color=MUTED)
    add_run(p, dates, size=9.8, color=MUTED)

    p = cell.add_paragraph()
    set_paragraph_spacing(p, after=3)
    add_run(p, role, bold=True, size=9.9, color=BLUE)

    add_bullets(cell, bullets)


def main():
    doc = Document()
    set_page(doc.sections[0])
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.2)

    hero = doc.add_table(rows=1, cols=1)
    hero.alignment = WD_TABLE_ALIGNMENT.CENTER
    hero.autofit = False
    hero.columns[0].width = Inches(7.0)
    set_table_borders(hero, color="20324D", sz="12")
    cell = hero.cell(0, 0)
    set_cell_shading(cell, "FFFFFF")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=2)
    add_run(p, "WESLEY SIMPLICIO", bold=True, size=20, color=INK)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=5)
    add_run(
        p,
        "Data/Evaluations Engineer | AI Agents | Evaluation Systems | Runtime Reliability",
        bold=True,
        size=10.2,
        color=BLUE,
    )

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=0, line=1.0)
    add_run(
        p,
        "Guarulhos, Sao Paulo, Brazil  |  +55 11 94669-4305  |  wesleysimplicio@live.com",
        size=9.2,
        color=MUTED,
    )

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=0, line=1.0)
    add_run(
        p,
        "GitHub: github.com/Wesleysimplicio  |  LinkedIn: br.linkedin.com/in/wesleysimplicio",
        size=9.2,
        color=MUTED,
    )

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(4.8)
    table.columns[1].width = Inches(2.2)
    set_table_borders(table, color="FFFFFF", sz="0")

    left = table.cell(0, 0)
    right = table.cell(0, 1)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    # LEFT COLUMN
    add_section_title(left, "Professional Summary")
    add_body(
        left,
        "Senior Full Stack Engineer with 10+ years of experience building production systems across finance, healthcare, retail, and enterprise software. Strong background in distributed systems, cloud infrastructure, automated testing, CI/CD, messaging, and software architecture for reliable and scalable applications.",
    )
    add_body(
        left,
        "Over the last year, I have worked deeply with AI agent engineering, repository-scale automation, evaluation-oriented workflows, and runtime tooling. My focus includes multi-agent orchestration, plan-first execution, validation loops, controlled write scopes, architecture-aware code generation, and operational guardrails across Claude Code, Codex, Cursor, GitHub Copilot, Hermes Agent, and OpenClaw.",
    )
    add_body(
        left,
        "I am especially interested in evaluation systems, failure analysis, runtime predictability, automated and human review loops, and the tooling needed to make agent behavior measurable and improvable at scale.",
    )

    add_section_title(left, "Evaluation and Agent Engineering Highlights")
    box = left.add_table(rows=1, cols=1)
    box.autofit = False
    box.columns[0].width = Inches(4.55)
    set_table_borders(box, color="CDDDFF", sz="8")
    bcell = box.cell(0, 0)
    set_cell_shading(bcell, "E9F1FF")
    p = bcell.paragraphs[0]
    set_paragraph_spacing(p, after=2)
    add_run(p, "Hermes Agent Runtime and Quality Improvements", bold=True, size=10.0, color=INK)
    p = bcell.add_paragraph()
    set_paragraph_spacing(p, after=0, line=1.08)
    add_run(
        p,
        "Contributed directly to Hermes Agent runtime quality, CI structure, and prompt assembly behavior, including work that improves determinism, stability, and future cache-aware execution.",
        size=9.6,
        color=INK,
    )

    add_bullets(
        left,
        [
            "Active contributor to the Hermes Agent ecosystem with more than 152 PRs opened in 24 hours, spanning bug fixes, runtime quality, workflow improvements, and test/CI infrastructure.",
            "Designed and implemented a provider-agnostic stable context prefix builder for Hermes Agent with deterministic ordering, stable hashing, structured segment modeling, and XML-like rendering.",
            "Built and maintained AI-friendly starter structures for multi-agent project bootstrap, standardized execution, documentation, and architecture-aware onboarding.",
            "Worked with repository-wide changes, controlled automation, validation-first workflows, read/write constraints, and repeatable instruction layers to reduce regressions and improve delivery quality.",
            "Standardized agent behavior across multiple CLIs and model providers to improve repeatability, safety, and operational consistency.",
        ],
    )

    add_section_title(left, "Selected Open Source Work")
    p = left.add_paragraph()
    set_paragraph_spacing(p, before=2, after=2)
    add_run(p, "Hermes Agent Contributions", bold=True, size=10.1, color=INK)
    add_bullets(
        left,
        [
            "PR #23479 - feat(runtime): add stable context prefix builder. Implemented a provider-agnostic StableContextBuilder to support stable and dynamic prompt assembly, future prefix caching, and more predictable runtime behavior.",
            "151 PRs created for Hermes Agent in 24 hours. Contributed a high-volume stream of fixes and improvements across runtime quality, workflow consistency, validation, and developer experience.",
        ],
    )
    p = left.add_paragraph()
    set_paragraph_spacing(p, before=2, after=2)
    add_run(p, "Agentic Starter", bold=True, size=10.1, color=INK)
    add_body(
        left,
        "AI-friendly, stack-neutral starter pack for multi-agent project bootstrap. Designed to standardize handoff, execution flow, documentation structure, instruction layers, and controlled agent operation across Claude Code, Codex, Copilot, Cursor, Hermes Agent, OpenClaw, and similar environments.",
    )
    p = left.add_paragraph()
    set_paragraph_spacing(p, before=2, after=2)
    add_run(p, "Other AI Projects", bold=True, size=10.1, color=INK)
    add_bullets(
        left,
        [
            "PiAPI-Skills / WaveSpeedAI-Skills - reusable skill packs with support for 700+ models.",
            "SendSprint - structured multi-agent sprint automation workflow.",
        ],
    )

    add_section_title(left, "Professional Experience")
    add_exp_entry(
        left,
        "Banco Fibra",
        "Oct 2025 - Feb 2026",
        "Senior Full Stack .NET / React Developer",
        [
            "Worked on modernization efforts migrating legacy systems into microservices and micro frontend architectures.",
            "Delivered improvements and fixes in .NET and React applications while preserving architecture and operational reliability.",
            "Used AI-assisted workflows to coordinate frontend and backend delivery under controlled execution constraints.",
        ],
    )
    add_exp_entry(
        left,
        "EY",
        "Jan 2025 - Oct 2025",
        "Senior Full Stack Developer (.NET Core 9 / Angular 19)",
        [
            "Improved internal workflow and enterprise process systems using .NET Core, Angular, Azure, and Entity Framework.",
            "Worked with CI/CD, Git Flow, Scrum, and automated testing in process-heavy enterprise environments.",
        ],
    )
    add_exp_entry(
        left,
        "Banco BMG",
        "Jan 2024 - Dec 2024",
        "Senior Full Stack Developer (.NET Core 8 / Angular)",
        [
            "Built and improved PIX-related systems aligned with Brazilian Central Bank requirements.",
            "Worked with AWS, RabbitMQ, Kafka, Entity Framework, and Hexagonal Architecture on reliable regulated systems.",
        ],
    )
    add_exp_entry(
        left,
        "IOB",
        "Jan 2022 - Dec 2023",
        "Senior Full Stack Developer (.NET Core 8 / React / Angular)",
        [
            "Maintained and evolved product features with strong attention to process improvement, stability, and scalable delivery.",
            "Worked with AWS, SQS, CI/CD, DDD, and Hexagonal Architecture.",
        ],
    )
    add_exp_entry(
        left,
        "Earlier Experience",
        "2016 - 2021",
        "Fast Shop, United Health Group / Amil, and TIVIT",
        [
            "Built enterprise systems using .NET, Angular, React, Azure Functions, Azure Service Bus, MongoDB, PostgreSQL, SQL Server, and microservices.",
            "Integrated enterprise systems with SAP and Microsiga while maintaining operational reliability across corporate environments.",
        ],
    )

    # RIGHT COLUMN
    add_section_title(right, "Best Match for Data/Evals Work")
    add_bullets(
        right,
        [
            "Strong experience with automated validation, CI pipelines, test reliability, and regression-resistant workflows.",
            "Hands-on work with failure reduction, deterministic runtime behavior, and measurable quality improvements in agent systems.",
            "Comfortable designing processes that combine automation and human review.",
            "Track record of OSS contributions and shipping practical improvements in active AI-agent codebases.",
        ],
    )

    add_section_title(right, "Core Technical Strengths")
    strengths = [
        "Evaluation workflows and validation systems",
        "Failure analysis and runtime reliability",
        "Controlled automation and review loops",
        "Distributed systems and microservices",
        "RabbitMQ, Kafka, AWS SQS, Azure Service Bus",
        "AWS and Azure cloud infrastructure",
        "CI/CD and automated testing",
        "xUnit, NUnit, Moq, TestContainers",
        "C#, .NET Core, Entity Framework, Dapper, Web API",
        "Angular, React, Vue.js, JavaScript, HTML5, CSS3",
        "DDD, Hexagonal Architecture, event-driven systems",
        "SQL Server, PostgreSQL, MySQL, Oracle, MongoDB",
    ]
    add_bullets(right, strengths)

    add_section_title(right, "AI Ecosystem Exposure")
    add_body(
        right,
        "Hands-on with Claude, GPT, Gemini, Grok, Kimi, GLM, MiniMax, Hermes Agent, Codex, Cursor, GitHub Copilot, and OpenClaw, with a focus on operational quality, predictability, and agent workflow standardization.",
    )

    add_section_title(right, "Education")
    add_body(right, "Bachelor's Degree in Systems Analysis and Development\nENIAC - Completed Dec 2014")
    add_body(right, "Technical Degree in Information Technology\nETEC Horacio")
    add_body(right, "Additional coursework\nWeb Design - Senac")

    add_section_title(right, "Links")
    add_body(right, "GitHub: github.com/Wesleysimplicio")
    add_body(right, "LinkedIn: br.linkedin.com/in/wesleysimplicio")
    add_body(right, "Highlighted Hermes PR: github.com/NousResearch/hermes-agent/pull/23479")

    doc.save(OUT_DOCX)


if __name__ == "__main__":
    main()
